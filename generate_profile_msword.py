import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from utils.hyperlocal_macro_factory import MacroFeatureEngineV2


def create_element(name):
    return OxmlElement(name)


def set_cell_background(cell, fill_hex):
    """Safely injects an institutional background color into a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def add_horizontal_divider(doc, hex_color="1B365D"):
    """Inserts a styled, colored horizontal rule to separate report dimensions."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = create_element("w:pBdr")
    bottom = create_element("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")  # Thickness indicator
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def append_metric_section(doc, primary_color, name, result, spectrum_dict):
    """Helper to append structured metric descriptions with risk spectra."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(2)
    h_run = h.add_run(f"站 {name}")
    h_run.font.size = Pt(12)
    h_run.font.bold = True
    h_run.font.color.rgb = primary_color

    # Display Current Cohort Calculation Result
    res_p = doc.add_paragraph()
    res_p.paragraph_format.left_indent = Inches(0.25)
    res_p.paragraph_format.space_after = Pt(4)
    lbl = res_p.add_run("Current Cohort Result: ")
    lbl.font.bold = True
    val = res_p.add_run(result)
    val.font.bold = True
    val.font.color.rgb = primary_color

    # Build cleanly aligned spectrum bounds
    for tier, text in spectrum_dict.items():
        spec_p = doc.add_paragraph()
        spec_p.paragraph_format.left_indent = Inches(0.5)
        spec_p.paragraph_format.space_after = Pt(2)

        tier_run = spec_p.add_run(f"•  {tier}: ")
        tier_run.font.bold = True
        spec_p.add_run(text)


def generate_word_loan_dashboard(profile_data: dict, export_path: Path):
    """
    Transforms raw engine output into an enterprise-grade Microsoft Word (.docx)
    underwriting report with professional layout hierarchies and risk espectrums.
    """
    fips = profile_data.get("target_fips", "UNKNOWN")
    year = profile_data.get("target_year", "UNKNOWN")
    naics = profile_data.get("target_naics", "GENERAL")
    flag = profile_data.get("spatial_governance_flag", "UNKNOWN")

    # Extract raw metrics with fallback protections
    m_cushion = profile_data.get("macro_wealth_cushion", 0.0) or 0.0
    f_velocity = profile_data.get("filer_density_velocity", 0.0) or 0.0
    h_dependency = profile_data.get("household_dependency_ratio", 0.0) or 0.0
    l_friction = profile_data.get("labor_pool_structural_friction", 0.0) or 0.0
    w_diversification = profile_data.get("wage_diversification_index", 1.0) or 1.0
    w_disconnect = profile_data.get("wage_to_filer_disconnect_index", 0.0) or 0.0
    s_momentum = profile_data.get("state_coincident_momentum", 0.0) or 0.0
    y_spread = profile_data.get("sovereign_yield_spread", 0.25) or 0.25
    c_sentiment = profile_data.get("macro_consumer_sentiment", 70.0) or 70.0
    p_turnover = profile_data.get("state_private_turnover_rate", 0.0) or 0.0
    n_job_flow = profile_data.get("state_net_job_flow_count", 0) or 0
    m_saturation = profile_data.get("industry_market_saturation_lq", None)

    doc = Document()

    # Configure 1-inch standard corporate margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Palette Brand Definitions (Hex and RGB)
    PRIMARY_COLOR = RGBColor(27, 54, 93)  # Deep Dark Blue (#1B365D)
    SECONDARY_COLOR = RGBColor(92, 102, 112)  # Slate Gray (#5C6670)
    TEXT_COLOR = RGBColor(34, 34, 34)  # Charcoal Off-Black

    # Base Document Style Configuration
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_COLOR

    # --- DOCUMENT HEADER ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("🛸 ENTERPRISE UNDERWRITING RISK REPORT")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(12)
    sub_run = subtitle_p.add_run("MACROECONOMIC UNDERWRITING DATA ENGINE")
    sub_run.font.size = Pt(10)
    sub_run.font.bold = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    add_horizontal_divider(doc, "1B365D")

    # --- META INFORMATION BOX (Styled Callout Table) ---
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.autofit = False

    meta_data = [
        [f"COHORT TARGET FIPS: {fips}", f"VINTAGE YEAR: {year}"],
        [f"TARGET NAICS SECTOR: {naics}", f"SPATIAL GOVERNANCE FLAG: {flag}"],
    ]

    for row_idx, row_content in enumerate(meta_data):
        for col_idx, text in enumerate(row_content):
            cell = meta_table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_background(cell, "F4F6F8")  # Light gray background pill
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(4)
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    # =====================================================================
    # DIMENSION 1: PASSIVE WEALTH DEPTH & MIGRATION PROFILES (IRS SOI)
    # =====================================================================
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(18)
    h1_run = h1.add_run("1. PASSIVE WEALTH DEPTH & MIGRATION PROFILES (IRS SOI)")
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = PRIMARY_COLOR

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "Macro Wealth Cushion",
        f"{m_cushion:.4f}",
        {
            "High (≥ 0.120)": "Strong local liquidity cushion. High concentration of investment dividends and interest relative to labor wages, indicating local wealth insulation.",
            "Average (0.040 to 0.119)": "Standard balanced marketplace. Consumer spending supported primarily by active jobs with normal household savings reserves.",
            "Low (< 0.040)": "High-risk capital constraints. Community depends entirely on immediate payroll cycles; households lack passive liquidity cushions during stress.",
        },
    )

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "Filer Density Velocity",
        f"{f_velocity:+.4f} ({f_velocity*100:+.2f}%)",
        {
            "High (≥ +0.050)": "Rapid regional expansion. Net taxpayer population grew by 5%+ over rolling 5-year window, indicating strong business and home demand.",
            "Average (-0.020 to +0.049)": "Stable economic baseline. Normal household turnover and steady organic migration patterns.",
            "Low (< -0.020)": "Regional contraction / Capital flight. Taxpayers are leaving the market, threatening the local tax base and commercial asset values.",
        },
    )

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "Household Dependency Ratio",
        f"{h_dependency:.4f}",
        {
            "High (≥ 2.20)": "Demographically vulnerable. Large family sizes or high non-working dependent population relative to tax filers, tightening disposable household income.",
            "Average (1.60 to 2.19)": "Balanced family demographic footprint matching national baseline cost structures.",
            "Low (< 1.60)": "Favorable household cash flow profile. High concentration of single filers or dual-income households with low dependency constraints.",
        },
    )

    # =====================================================================
    # DIMENSION 2: LABOR MARKET DYNAMICS & FRICTION
    # =====================================================================
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(14)
    h2_run = h2.add_run("2. LABOR MARKET DYNAMICS & FRICTION (BLS LAUS / CENSUS QWI)")
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = PRIMARY_COLOR

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "Labor Pool Structural Friction",
        f"{l_friction:.4f}",
        {
            "High (≥ 0.060)": "Volatile workforce environment. Erratic labor force counts or extreme seasonal employment spikes that can interrupt corporate operations.",
            "Average (0.015 to 0.059)": "Stable labor pool layout. Predictable worker availability for consistent regional business hiring.",
            "Low (< 0.015)": "Flat workforce landscape. Indicates a stagnant or rigid local labor market, making it difficult for scaling firms to find and hire new staff.",
        },
    )

    if m_saturation is not None:
        append_metric_section(
            doc,
            PRIMARY_COLOR,
            "Industry Market Saturation (LQ)",
            f"{m_saturation:.4f}",
            {
                "High (≥ 1.25)": "Highly specialized exporter sector. The market has an intense cluster of this business type compared to the US baseline, making it sensitive to sector cycles.",
                "Average (0.75 to 1.24)": "Equilibrium cluster density. Local business volume matches standard domestic consumption patterns perfectly.",
                "Low (< 0.75)": "Underserved local market. The sector is underrepresented, signaling potential market entry space or weak regional consumer demand.",
            },
        )

    # =====================================================================
    # DIMENSION 3: ADVANCED REGIONAL STRUCTURE & DISCONNECT SPREADS
    # =====================================================================
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(14)
    h3_run = h3.add_run(
        "3. ADVANCED REGIONAL STRUCTURE & DISCONNECT SPREADS (BLS QCEW)"
    )
    h3_run.font.size = Pt(14)
    h3_run.font.bold = True
    h3_run.font.color.rgb = PRIMARY_COLOR

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "Wage Diversification Index (Inverse Payroll HHI)",
        f"{w_diversification:.4f}",
        {
            "High (≥ 0.920)": "Resilient, diverse regional economy. Total corporate payroll is spread smoothly across many distinct industries, insulating the market from single-sector crashes.",
            "Average (0.750 to 0.919)": "Normal industry balance. Typical industrial mix with mild leaning toward localized anchor fields.",
            "Low (< 0.750)": "Monopolized / Vulnerable 'Company Town.' Payroll concentration is dominated by a few corporate players, leaving local debt service vulnerable if that sector slips.",
        },
    )

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "Wage-to-Filer Disconnect Index",
        f"{w_disconnect:+.4f}",
        {
            "High (≥ +0.040)": "Widening structural gap. IRS resident taxpayer wage growth is outstripping local business job pay, signaling affluent commuters moving in or gentrification.",
            "Average (-0.039 to +0.039)": "Symmetric alignment. Wage growth inside local establishments matches local resident income tracking closely.",
            "Low (< -0.040)": "Commercial operational stress. Local business wages are outstripping resident income growth, indicating rising business overhead or local worker scarcity.",
        },
    )

    # =====================================================================
    # DIMENSION 4: SYSTEMIC SOVEREIGN & STATE SHOCK ANCHORS
    # =====================================================================
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(14)
    h4_run = h4.add_run(
        "4. SYSTEMIC SOVEREIGN & STATE SHOCK ANCHORS (FRED / CENSUS STATE)"
    )
    h4_run.font.size = Pt(14)
    h4_run.font.bold = True
    h4_run.font.color.rgb = PRIMARY_COLOR

    append_metric_section(
        doc,
        PRIMARY_COLOR,
        "State Coincident Momentum",
        f"{s_momentum:+.4f} ({s_momentum*100:+.2f}%)",
        {
            "High (≥ +0.035)": "High-velocity boom track. The state economy is expanding rapidly, providing strong macro tailwinds for local business revenues.",
            "Average (0.000 to +0.034)": "Steady, sustainable economic growth. Solid foundation for standard long-term underwriting horizons.",
            "Low (< 0.000)": "Systemic state-level recession. The regional economy is actively contracting, signaling elevated credit risks across all portfolios.",
        },
    )

    # Systemic Tail Risks Grid Table
    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    tail_table = doc.add_table(rows=4, cols=2)

    tail_data = [
        ["Sovereign Yield Spread (T10Y2Y)", f"{y_spread:+.2f}%"],
        ["Macro Consumer Sentiment (UMCSENT)", f"{c_sentiment:.1f}"],
        ["State Private Labor Turnover Rate", f"{p_turnover*100:.2f}%"],
        ["State Net Job Flow Count", f"{n_job_flow:,}"],
    ]

    for idx, (label, val_text) in enumerate(tail_data):
        c0 = tail_table.cell(idx, 0)
        c1 = tail_table.cell(idx, 1)
        c0.text = label
        c1.text = val_text
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        c0.paragraphs[0].runs[0].font.bold = True
        c1.paragraphs[0].runs[0].font.bold = True
        if idx % 2 == 0:
            set_cell_background(c0, "FAFAFA")
            set_cell_background(c1, "FAFAFA")

    # Footer Metadata and Legal Disclaimers
    add_horizontal_divider(doc, "5C6670")
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run(
        "CONFIDENTIAL INSTITUTIONAL USE ONLY — GENERATED VIA MACROFEATUREENGINEV2"
    )
    f_run.font.size = Pt(8)
    f_run.font.color.rgb = SECONDARY_COLOR
    f_run.font.italic = True

    # Save output asset path safely to the local filesystem
    doc.save(str(export_path))
    print(f"📄 Compliance Word Report successfully written to disk: {export_path}")


engine = MacroFeatureEngineV2
# Extract the clean, optimized macro parameters dictionary
raw_profile = engine.extract_comprehensive_macro_profile(
    county_fips="48085", loan_vintage_year=2016, naics_4d="4411"
)

# Compile and export directly into a Word document file on disk
output_file = Path("Deal_Folder_FIPS_48085_Vintage_2016.docx")
generate_word_loan_dashboard(raw_profile, output_file)
